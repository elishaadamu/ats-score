import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def configure_styles(doc, font_name, h1_size, h1_color, h2_size, h2_color, align_centered=False):
    # Configure Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = font_name
    style_normal.font.size = Pt(11)
    
    # Configure Heading 1
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = font_name
    style_h1.font.size = Pt(h1_size)
    style_h1.font.color.rgb = h1_color
    style_h1.font.bold = True
    if align_centered:
        style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    # Configure Heading 2
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = font_name
    style_h2.font.size = Pt(h2_size)
    style_h2.font.color.rgb = h2_color
    style_h2.font.bold = True

def main():
    out_dir = "public/templates"
    os.makedirs(out_dir, exist_ok=True)
    
    # ==========================
    # 1. CLASSIC TEMPLATE
    # ==========================
    doc1 = Document()
    configure_styles(doc1, 'Times New Roman', 24, RGBColor(0, 0, 0), 14, RGBColor(50, 50, 50), align_centered=True)
    doc1.add_heading("APPLICANT NAME", level=1)
    doc1.add_paragraph("Email: user@example.com | Phone: 555-555-5555 | LinkedIn: linkedin.com/in/user").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc1.add_heading("PROFESSIONAL SUMMARY", level=2)
    doc1.add_paragraph("Classic structured summary placed here.")
    doc1.save(os.path.join(out_dir, "Template_1.docx"))

    # ==========================
    # 2. MODERN TEMPLATE
    # ==========================
    doc2 = Document()
    configure_styles(doc2, 'Arial', 26, RGBColor(0, 102, 204), 14, RGBColor(0, 102, 204), align_centered=False)
    doc2.add_heading("APPLICANT NAME", level=1)
    doc2.add_paragraph("Email: user@example.com | Phone: 555-555-5555\nPortfolio: portfolio.example.com")
    doc2.add_heading("PROFESSIONAL SUMMARY", level=2)
    doc2.add_paragraph("Modern left-aligned text with bright blue accents.")
    doc2.save(os.path.join(out_dir, "Template_2.docx"))

    # ==========================
    # 3. EXECUTIVE TEMPLATE
    # ==========================
    doc3 = Document()
    configure_styles(doc3, 'Garamond', 28, RGBColor(0, 51, 102), 16, RGBColor(0, 0, 0), align_centered=True)
    doc3.add_heading("APPLICANT NAME", level=1)
    doc3.add_paragraph("user@example.com | 555-555-5555").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc3.add_heading("EXECUTIVE SUMMARY", level=2)
    doc3.add_paragraph("Sophisticated layout using Garamond, elegant center alignments and dark navy colors.")
    doc3.save(os.path.join(out_dir, "Template_3.docx"))

    print("Templates 1, 2, 3 generated.")

if __name__ == "__main__":
    main()
